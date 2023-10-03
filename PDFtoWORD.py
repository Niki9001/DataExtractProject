import pdfplumber
from docx import Document

for j in range(1,187):
    # 创建一个新的Word文档
    doc = Document()
    #pdf存储地址
    url = f'H:\Training_profile\Profile ({j}).pdf'

# 打开PDF文件
    with pdfplumber.open(url) as pdf:
    # 遍历每一页
        for i, page in enumerate(pdf.pages):
            # 获取整个页面的尺寸
            page_width = page.width
            page_height = page.height

            # 定义左侧区域：(x0, y0, x1, y1)
            left_box = (0, 0, page_width / 2.8, page_height)

            # 定义右侧区域：(x0, y0, x1, y1)
            right_box = (185,0,page_width , page_height)

            # 裁剪左侧区域并提取文本
            left_crop = page.crop(left_box)
            left_text = left_crop.extract_text()

            # 裁剪右侧区域并提取文本
            right_crop = page.crop(right_box)
            right_text = right_crop.extract_text()

            # 将提取的文本添加到Word文档中
            doc.add_heading(f"Page {i + 1}", level=0)
            doc.add_heading("Left Side", level=1)
            doc.add_paragraph(left_text)
            doc.add_heading("Right Side", level=1)
            doc.add_paragraph(right_text)
            # 保存Word文档
            doc.save(f"D:\PycharmProjects\pythonProject4\extracted_profile ({j}).docx")


        # 打印或保存文本
#        print(f"Left text from page {i + 1}:\n{left_text}")
#        print('------------------------------------------')
#        print(f"Right text from page {i + 1}:\n{right_text}")
